import requests
from bs4 import BeautifulSoup
from openai import OpenAI
from docx import Document
from docx.shared import Pt

# Replace 'your-api-key' with your actual OpenAI API key
api_key = "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"
client = OpenAI(api_key=api_key)

def get_patent_info(patent_number):
    url = f"https://patents.google.com/patent/{patent_number}/en"
    response = requests.get(url)
    soup = BeautifulSoup(response.content, "html.parser")
    # Extracting the required fields
    title = soup.find("meta", {"name": "DC.title"}).get("content", "N/A")
    title = title.replace("\n", "")
    assignee = (
        soup.find("dd", {"itemprop": "assigneeCurrent"}).get_text(strip=True)
        if soup.find("dd", {"itemprop": "assigneeCurrent"})
        else "N/A"
    )
    application_date = soup.find("meta", {"name": "DC.date"}).get("content", "N/A")
    status = (
        soup.find("span", {"itemprop": "status"}).get_text(strip=True)
        if soup.find("span", {"itemprop": "status"})
        else "N/A"
    )
    claims_abstract = (
        soup.find("section", {"itemprop": "claims"}).get_text(strip=True)
        + "\n"
        + soup.find("section", {"itemprop": "abstract"}).get_text(strip=True)
    )
    return {
        "Publication Number": patent_number,
        "Title": title,
        "Assignee": assignee,
        "Application Date": application_date,
        "Status": status,
        "Claims and Abstract": claims_abstract,
    }

def generate_summary(claims_abstract):
    prompt = f"Generate a summary for the following claims and abstract in around 100 words:\n\n{claims_abstract}"
    response = client.chat.completions.create(
        model="gpt-4",  # Use the appropriate model for your use case
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt},
        ],
        max_tokens=100,
        temperature=0.7,
    )
    summary = response.choices[0].message.content.strip()
    return summary

def write_to_word(patent_info, summary, document):
    # Function to add a bold text run followed by regular text in the same paragraph
    def add_bold_text(doc, bold_text, normal_text):
        p = doc.add_paragraph()
        run = p.add_run(bold_text)
        run.bold = True
        run.font.size = Pt(12)
        p.add_run(normal_text)

    # Adding the patent information to the Word document
    add_bold_text(document, "Publication Number: ", patent_info["Publication Number"])
    add_bold_text(document, "Title: ", patent_info["Title"])
    add_bold_text(document, "Assignee: ", patent_info["Assignee"])
    add_bold_text(document, "Application Date: ", patent_info["Application Date"])
    add_bold_text(document, "Status: ", patent_info["Status"])
    add_bold_text(document, "Summary: ", summary)
    document.add_paragraph(" " * 2)

# Input: List of patent numbers separated by a newline character
patent_numbers = """
US20230140553A1
US11457554B2
US11727119B2
US10963231B1
US11367008B2
US11816244B2
US20220318577A1
US20230342392A1
US11983609B2
US20230123231A1
US20230245008A1
US11410111B1
US20220027793A1
US9576262B2
US20210374866A1
US20210065305A1
US20200294073A1
US20230011954A1
US11223546B2
US20210216911A1
"""
# Create a new Word document
document = Document()
# Process each patent number and write to the Word document
for patent_number in patent_numbers.splitlines():
    if patent_number.strip():
        patent_info = get_patent_info(patent_number.strip())
        summary = generate_summary(patent_info["Claims and Abstract"])
        write_to_word(patent_info, summary, document)
# Save the Word document
document.save("patent_info.docx")
print("Word document created successfully!")